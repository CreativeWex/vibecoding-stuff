#!/usr/bin/env python3
"""Build БД_Лабораторная_работа.docx from steps.json (python-docx)."""

from __future__ import annotations

import argparse
import json
import logging
import sys
from pathlib import Path

_SCRIPT_DIR = Path(__file__).resolve().parent
if str(_SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPT_DIR))

from docx import Document
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.shared import Pt

from logging_config import setup_logging

log = logging.getLogger(__name__)

# Текст, который в старых steps.json означал «параграф не нужен».
_NO_NOTES_SENTINELS = frozenset(
    {
        "не требуется",
        "не требуется.",
        "n/a",
        "—",
        "-",
    }
)


def _should_emit_notes(notes_raw: object) -> bool:
    """Параграф «Дополнительные инструкции» только при непустом осмысленном тексте."""
    if notes_raw is None or not isinstance(notes_raw, str):
        return False
    text = notes_raw.strip()
    if not text:
        return False
    if text.casefold() in _NO_NOTES_SENTINELS:
        return False
    return True


def _append_notes_body(paragraph, notes: str) -> None:
    """Многострочные notes: переносы \\n → мягкий перевод строки в Word."""
    lines = notes.split("\n")
    for j, line in enumerate(lines):
        if j > 0:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        paragraph.add_run(line)


def build_lab_docx(steps_data: dict, output_path: Path) -> None:
    title = steps_data.get("title") or "Лабораторная работа по базам данных"
    steps = steps_data.get("steps") or []
    if not steps:
        raise ValueError("steps.json: empty or missing 'steps' array")

    doc = Document()
    doc.add_heading(title, level=0)

    for i, step in enumerate(steps, start=1):
        doc.add_heading(f"Шаг {i}", level=1)
        desc = step.get("description", "").strip() or "—"
        sql = step.get("sql", "")
        notes_raw: object = step.get("notes")

        p1 = doc.add_paragraph()
        p1.add_run("Описание: ").bold = True
        p1.add_run(desc)

        p2 = doc.add_paragraph()
        p2.add_run("SQL: ").bold = True
        for line in sql.splitlines() or [""]:
            r = p2.add_run(line + "\n")
            r.font.name = "Courier New"
            r.font.size = Pt(11)
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        if _should_emit_notes(notes_raw):
            notes = str(notes_raw).strip()
            p3 = doc.add_paragraph()
            p3.add_run("Дополнительные инструкции: ").bold = True
            _append_notes_body(p3, notes)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    log.info("Written docx: %s (%s steps)", output_path, len(steps))


def main() -> int:
    ap = argparse.ArgumentParser(description="Build lab report docx from steps.json")
    ap.add_argument("--steps", type=Path, required=True, help="Path to steps.json")
    ap.add_argument("--output", type=Path, required=True, help="Output .docx path")
    ap.add_argument("-v", "--verbose", action="store_true")
    args = ap.parse_args()
    setup_logging(args.verbose)

    try:
        raw = args.steps.read_text(encoding="utf-8", errors="replace")
        data = json.loads(raw)
    except Exception as e:
        log.error("Cannot read steps.json: %s", e)
        return 1

    try:
        build_lab_docx(data, args.output)
    except Exception as e:
        log.error("%s", e)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
