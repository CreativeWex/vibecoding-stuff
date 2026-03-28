#!/usr/bin/env python3
"""Build an empty GOST-styled Word template (styles + page setup)."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Mm, Pt


def _set_heading_styles(doc: Document) -> None:
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.italic = False
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.first_line_indent = Cm(0)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h1.paragraph_format.line_spacing = 1.5

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.italic = False
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.first_line_indent = Cm(0)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h2.paragraph_format.line_spacing = 1.5

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.italic = True
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.first_line_indent = Cm(0)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h3.paragraph_format.line_spacing = 1.5

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.5


def _set_margins(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_height = Mm(297)
    sec.page_width = Mm(210)
    sec.left_margin = Mm(25)
    sec.right_margin = Mm(10)
    sec.top_margin = Mm(20)
    sec.bottom_margin = Mm(20)


def build_template(out_path: Path) -> None:
    doc = Document()
    _set_margins(doc)
    _set_heading_styles(doc)
    doc.add_paragraph("Заголовок документа (Heading 1)", style="Heading 1")
    doc.add_paragraph(
        "Пример основного текста: абзац с интервалом 1,5 и красной строкой 1,25 см, "
        "выравнивание по ширине.",
        style="Normal",
    )
    doc.save(out_path)


def main() -> None:
    p = argparse.ArgumentParser(description="Create GOST template .docx")
    p.add_argument(
        "-o",
        "--output",
        type=Path,
        default=Path(__file__).resolve().parent.parent / "template_gost.docx",
        help="Output path for template_gost.docx",
    )
    args = p.parse_args()
    args.output.parent.mkdir(parents=True, exist_ok=True)
    build_template(args.output)
    print(f"Template written: {args.output}")


if __name__ == "__main__":
    main()
