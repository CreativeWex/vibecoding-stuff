#!/usr/bin/env python3
"""
Apply typography and page layout aligned with common methodology for ГОСТ Р 7.0.11-2008
and typical university lab/thesis guides (.docx via python-docx).

Limitations: .doc is not supported. TOC/PAGE fields must be updated in Word.
"""

from __future__ import annotations

import argparse
import re
import shutil
import sys
import time
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt
from docx.text.paragraph import Paragraph

PAGE_W_MM, PAGE_H_MM = 210, 297

# Presets: "methodology" — частые сводки (25/10/20/20); "legacy" — прежние значения skill (30/15/20/20)
MARGIN_PRESETS: dict[str, dict[str, int]] = {
    "methodology": {"left": 25, "right": 10, "top": 20, "bottom": 20},
    "legacy": {"left": 30, "right": 15, "top": 20, "bottom": 20},
}

FIG_RE = re.compile(
    r"^(?P<label>Рисунок|Рис\.|Figure)\s*(?P<num>\d+)\s*(?P<sep>[-—–]\s*)?(?P<rest>.*)$",
    re.IGNORECASE | re.UNICODE,
)
TAB_RE = re.compile(
    r"^(?P<label>Таблица|Табл\.|Table)\s*(?P<num>\d+)\s*(?P<sep>[-—–]\s*)?(?P<rest>.*)$",
    re.IGNORECASE | re.UNICODE,
)

# Короткие типовые заголовки разделов отчёта / курсовой (без стилей Heading в исходнике)
_LAB_H2_PATTERNS = (
    r"^цель(\s+работы)?\.?$",
    r"^задание\.?$",
    r"^ход\s+работы\.?$",
    r"^теоретическая\s+часть\.?$",
    r"^практическая\s+часть\.?$",
    r"^выводы?\.?$",
    r"^контрольные\s+вопросы\.?$",
    r"^список\s+литературы\.?$",
    r"^введение\.?$",
    r"^заключение\.?$",
    r"^приложени[ея]\s*[а-я]?\)?\.?$",
    r"^оборудование\.?$",
    r"^результаты\.?$",
    r"^обсуждение\.?$",
    r"^методика\.?$",
    r"^эксперимент(альная\s+часть)?\.?$",
)
_LAB_H2_COMPILED = [re.compile(p, re.IGNORECASE | re.UNICODE) for p in _LAB_H2_PATTERNS]

# Длинное тире в подписях (ГОСТ-подобная практика)
CAP_DASH = " — "


def _is_figure_caption_text(raw: str) -> bool:
    t = raw.strip()
    return bool(FIG_RE.match(t)) or bool(re.match(r"^(Рисунок|Рис\.|Figure)\b", t, re.I))


def _is_table_caption_text(raw: str) -> bool:
    t = raw.strip()
    return bool(TAB_RE.match(t)) or bool(re.match(r"^(Таблица|Табл\.|Table)\b", t, re.I))


def _finalize_caption_title(text: str, fallback: str) -> str:
    t = re.sub(r"\s+", " ", text.strip())
    while t.endswith("."):
        t = t[:-1].strip()
    if not t:
        return fallback
    if len(t) > 100:
        t = t[:97].rstrip() + "…"
    return t[0].upper() + t[1:] if len(t) > 1 else t.upper()


def _snippet_for_caption(text: str, max_len: int = 90) -> str:
    t = re.sub(r"\s+", " ", text.strip())
    one = t.split(".")[0].strip()
    if len(one) > max_len:
        one = one[: max_len - 1].rstrip() + "…"
    return one


def _infer_figure_title_from_context(doc: Document, drawing_el) -> str:
    """Краткое название по соседним абзацам тела документа (без LLM)."""
    fallback = "Иллюстрация к работе"
    prev = drawing_el.getprevious()
    nxt = drawing_el.getnext()
    candidates: list[str] = []

    def _para_text(el) -> str:
        if el is None or el.tag != qn("w:p"):
            return ""
        return Paragraph(el, doc._body).text.strip()

    if prev is not None:
        t = _para_text(prev)
        if t and not _is_figure_caption_text(t) and not _is_table_caption_text(t):
            if not _lab_section_heading_level(t):
                candidates.append(_snippet_for_caption(t))

    if nxt is not None:
        t = _para_text(nxt)
        if t and not _is_figure_caption_text(t) and not _is_table_caption_text(t):
            if not _lab_section_heading_level(t):
                candidates.append(_snippet_for_caption(t))

    for c in candidates:
        if len(c) >= 3:
            return _finalize_caption_title(c, fallback)
    return fallback


def _table_header_guess(table: Table) -> str:
    if not table.rows:
        return ""
    cells = [c.text.strip() for c in table.rows[0].cells if c.text.strip()]
    if not cells:
        return ""
    joined = ", ".join(cells[:5])
    return joined if len(joined) <= 120 else joined[:117] + "…"


def _infer_table_title(doc: Document, table: Table) -> str:
    fallback = "Данные таблицы"
    prev_el = table._tbl.getprevious()
    if prev_el is not None and prev_el.tag == qn("w:p"):
        t = Paragraph(prev_el, doc._body).text.strip()
        if t and not _is_table_caption_text(t) and len(t) < 150:
            return _finalize_caption_title(_snippet_for_caption(t, 80), fallback)
    hg = _table_header_guess(table)
    if hg:
        inner = _finalize_caption_title(hg.split(",")[0].strip(), "Таблица")
        return f"Структура таблицы «{inner}»"
    return fallback


def _table_has_caption_above(doc: Document, table: Table) -> bool:
    prev = table._tbl.getprevious()
    if prev is None or prev.tag != qn("w:p"):
        return False
    return _is_table_caption_text(Paragraph(prev, doc._body).text)


def _insert_body_paragraph_before(doc: Document, element, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    element.addprevious(new_p)
    para = Paragraph(new_p, doc._body)
    para.add_run(text)
    return para


def _ensure_table_captions(doc: Document) -> int:
    """Подпись над таблицей: «Таблица N — …» (номер позже сквозной перенумерацией)."""
    added = 0
    for el in list(doc.element.body):
        if el.tag != qn("w:tbl"):
            continue
        table = Table(el, doc._body)
        if _table_has_caption_above(doc, table):
            continue
        title = _infer_table_title(doc, table)
        line = f"Таблица 1{CAP_DASH}{title}"
        para = _insert_body_paragraph_before(doc, table._tbl, line)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Cm(0)
        added += 1
    return added


def _ensure_figure_captions(doc: Document) -> int:
    """Подпись под рисунком; текст названия — из контекста."""
    added = 0
    changed = True
    while changed:
        changed = False
        for el in list(doc.element.body):
            if el.tag != qn("w:p"):
                continue
            p = Paragraph(el, doc._body)
            if not _paragraph_has_drawing(p):
                continue
            nxt = el.getnext()
            if nxt is not None and nxt.tag == qn("w:p"):
                if _is_figure_caption_text(Paragraph(nxt, doc._body).text):
                    continue
            title = _infer_figure_title_from_context(doc, el)
            line = f"Рисунок 1{CAP_DASH}{title}"
            para = _insert_body_paragraph_after(doc, el, line)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            added += 1
            changed = True
            break
    return added


def _insert_body_paragraph_after(doc: Document, p_el, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    p_el.addnext(new_p)
    para = Paragraph(new_p, doc._body)
    para.add_run(text)
    return para


def _add_page_field(run) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_end)


def _clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def _paragraph_has_drawing(paragraph) -> bool:
    if not paragraph.runs:
        return False
    return any("w:drawing" in run._element.xml or "w:pict" in run._element.xml for run in paragraph.runs)


def _setup_footer_page_numbers(doc: Document) -> None:
    """Номер снизу по центру со 2-й страницы (первая страница — пустой нижний колонтитул)."""
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    fp = sec.first_page_footer
    if fp.paragraphs:
        _clear_paragraph(fp.paragraphs[0])
    else:
        fp.add_paragraph()
    footer = sec.footer
    if not footer.paragraphs:
        footer.add_paragraph()
    p = footer.paragraphs[0]
    _clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    _add_page_field(run)


def _setup_header_page_numbers(doc: Document) -> None:
    """Номер сверху по центру со 2-й страницы; первый лист — отдельный верхний колонтитул (не затираем текст)."""
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    hdr = sec.header
    if not hdr.paragraphs:
        hdr.add_paragraph()
    p = hdr.paragraphs[0]
    _clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    _add_page_field(run)


def _apply_section_margins(doc: Document, margins_mm: dict[str, int]) -> None:
    for sec in doc.sections:
        sec.page_width = Mm(PAGE_W_MM)
        sec.page_height = Mm(PAGE_H_MM)
        sec.left_margin = Mm(margins_mm["left"])
        sec.right_margin = Mm(margins_mm["right"])
        sec.top_margin = Mm(margins_mm["top"])
        sec.bottom_margin = Mm(margins_mm["bottom"])


def _style_document_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.5

    for name, cfg in (
        (
            "Heading 1",
            {
                "size": 16,
                "bold": True,
                "italic": False,
                "align": WD_ALIGN_PARAGRAPH.CENTER,
                "before": 24,
                "after": 12,
            },
        ),
        (
            "Heading 2",
            {
                "size": 14,
                "bold": True,
                "italic": False,
                "align": WD_ALIGN_PARAGRAPH.LEFT,
                "before": 12,
                "after": 6,
            },
        ),
        (
            "Heading 3",
            {
                "size": 14,
                "bold": True,
                "italic": True,
                "align": WD_ALIGN_PARAGRAPH.LEFT,
                "before": 6,
                "after": 6,
            },
        ),
    ):
        if name not in doc.styles:
            continue
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(cfg["size"])
        st.font.bold = cfg["bold"]
        st.font.italic = cfg["italic"]
        st.paragraph_format.alignment = cfg["align"]
        st.paragraph_format.space_before = Pt(cfg["before"])
        st.paragraph_format.space_after = Pt(cfg["after"])
        st.paragraph_format.first_line_indent = Cm(0)
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        st.paragraph_format.line_spacing = 1.5

    for cap_name in ("Caption", "Подпись"):
        if cap_name in doc.styles:
            c = doc.styles[cap_name]
            c.font.name = "Times New Roman"
            c.font.size = Pt(12)


def _is_heading_style(name: str | None) -> bool:
    if not name:
        return False
    n = name.lower()
    return n.startswith("heading") or "заголовок" in n


def _heading_level(name: str | None) -> int | None:
    if not name:
        return None
    n = name.lower()
    if "heading 1" in n or "заголовок 1" in n:
        return 1
    if "heading 2" in n or "заголовок 2" in n:
        return 2
    if "heading 3" in n or "заголовок 3" in n:
        return 3
    return None


def _lab_section_heading_level(text: str) -> int | None:
    t = text.strip()
    if not t or len(t) > 120 or "\n" in t:
        return None
    for rx in _LAB_H2_COMPILED:
        if rx.match(t):
            return 2
    return None


def _snapshot_centered_paragraph_ids(doc: Document) -> set[int]:
    ids: set[int] = set()
    for p in _iter_all_paragraphs(doc):
        if p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            ids.add(id(p))
    return ids


def _apply_runs_font(paragraph, size_pt: float, bold: bool | None = None, italic: bool | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic


def _format_paragraph_body(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    _apply_runs_font(paragraph, 14)


def _format_title_block(paragraph) -> None:
    """Титул / центрированные блоки: без красной строки, по центру."""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    _apply_runs_font(paragraph, 14)


def _format_heading(paragraph, level: int) -> None:
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.5
    if level == 1:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(24)
        paragraph.paragraph_format.space_after = Pt(12)
        _apply_runs_font(paragraph, 16, bold=True, italic=False)
    elif level == 2:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        _apply_runs_font(paragraph, 14, bold=True, italic=False)
    else:
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        _apply_runs_font(paragraph, 14, bold=True, italic=True)


def _format_figure_caption(paragraph) -> None:
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.5
    _apply_runs_font(paragraph, 12, bold=False, italic=True)


def _format_table_caption(paragraph) -> None:
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.5
    _apply_runs_font(paragraph, 12, bold=False, italic=False)


def _looks_figure_caption(text: str) -> bool:
    t = text.strip()
    return bool(FIG_RE.match(t)) or t.lower().startswith("figure ") or bool(re.match(r"^рис\.\s*\d+", t, re.I))


def _looks_table_caption(text: str) -> bool:
    t = text.strip()
    return bool(TAB_RE.match(t)) or t.lower().startswith("table ") or bool(re.match(r"^табл\.\s*\d+", t, re.I))


def _process_paragraph(
    paragraph,
    *,
    centered_ids: set[int],
    preserve_center: bool,
    lab_heading_heuristic: bool,
    toc_field_id: int | None = None,
) -> None:
    if toc_field_id is not None and id(paragraph) == toc_field_id:
        _format_toc_field_paragraph(paragraph)
        return
    style_name = paragraph.style.name if paragraph.style else ""
    text = paragraph.text.strip()
    lvl = _heading_level(style_name)
    if lvl is not None:
        _format_heading(paragraph, lvl)
        return
    if _is_heading_style(style_name) and not lvl:
        _format_heading(paragraph, 1)
        return
    if "caption" in style_name.lower() or "подпись" in style_name.lower():
        if _looks_figure_caption(text):
            _format_figure_caption(paragraph)
        elif _looks_table_caption(text):
            _format_table_caption(paragraph)
        else:
            _format_table_caption(paragraph)
        return
    if _looks_figure_caption(text):
        _format_figure_caption(paragraph)
        return
    if _looks_table_caption(text):
        _format_table_caption(paragraph)
        return
    if "footnote" in style_name.lower() or "сноска" in style_name.lower():
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _apply_runs_font(paragraph, 10)
        return
    if preserve_center and id(paragraph) in centered_ids:
        _format_title_block(paragraph)
        return
    if lab_heading_heuristic:
        h = _lab_section_heading_level(text)
        if h is not None:
            _format_heading(paragraph, h)
            return
    _format_paragraph_body(paragraph)


def _iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _center_tables(doc: Document) -> None:
    for table in doc.tables:
        try:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        except Exception:
            pass
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(6)


def _apply_tnr_to_table_cells(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Times New Roman"


def _apply_tnr_to_body_drawing_paragraphs(doc: Document) -> None:
    for el in doc.element.body:
        if el.tag != qn("w:p"):
            continue
        p = Paragraph(el, doc._body)
        if not _paragraph_has_drawing(p):
            continue
        for run in p.runs:
            run.font.name = "Times New Roman"


def _center_image_paragraphs(doc: Document) -> None:
    for p in doc.paragraphs:
        if not p.runs:
            continue
        has_drawing = _paragraph_has_drawing(p)
        if has_drawing:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)


def _renumber_captions(doc: Document) -> tuple[int, int]:
    """Сквозная нумерация подписей в порядке следования в теле документа (не внутри ячеек)."""
    fig_n = 0
    tab_n = 0
    for el in doc.element.body:
        if el.tag != qn("w:p"):
            continue
        p = Paragraph(el, doc._body)
        raw = p.text.strip()
        if not raw:
            continue
        m = FIG_RE.match(raw)
        if m:
            fig_n += 1
            rest = (m.group("rest") or "").strip()
            sep = CAP_DASH if rest else ""
            label = m.group("label")
            if label.lower() in ("figure",):
                new_t = f"Figure {fig_n}{sep}{rest}"
            else:
                new_t = f"Рисунок {fig_n}{sep}{rest}"
            _replace_paragraph_text(p, new_t)
            continue
        m2 = TAB_RE.match(raw)
        if m2:
            tab_n += 1
            rest = (m2.group("rest") or "").strip()
            sep = CAP_DASH if rest else ""
            label = m2.group("label")
            if label.lower() in ("table",):
                new_t = f"Table {tab_n}{sep}{rest}"
            else:
                new_t = f"Таблица {tab_n}{sep}{rest}"
            _replace_paragraph_text(p, new_t)
    return fig_n, tab_n


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(new_text)


def _count_style_headings(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        sn = p.style.name if p.style else ""
        if _heading_level(sn) is not None:
            n += 1
        elif _is_heading_style(sn):
            n += 1
    return n


def _count_heuristic_headings_body(doc: Document, centered_ids: set[int]) -> int:
    n = 0
    for p in doc.paragraphs:
        if id(p) in centered_ids:
            continue
        if _lab_section_heading_level(p.text):
            n += 1
    return n


def _effective_heading_count_for_toc(doc: Document, centered_ids: set[int]) -> int:
    return _count_style_headings(doc) + _count_heuristic_headings_body(doc, centered_ids)


def _insert_toc_paragraph(doc: Document) -> Paragraph:
    """Возвращает абзац с полем TOC (без красной строки при последующей обработке)."""
    first = doc.paragraphs[0]
    p = first.insert_paragraph_before()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_end)
    p2 = p.insert_paragraph_before()
    p2.add_run("Содержание").bold = True
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _format_toc_field_paragraph(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.5
    _apply_runs_font(paragraph, 14)


def format_docx(
    input_path: Path,
    output_path: Path,
    *,
    backup: bool = False,
    insert_toc: bool | None = None,
    toc_min_headings: int = 3,
    renumber: bool = True,
    progress: bool = False,
    margin_preset: str = "methodology",
    margins_mm: dict[str, int] | None = None,
    page_number: str = "header",
    preserve_center: bool = True,
    lab_heading_heuristic: bool = True,
    infer_captions: bool = True,
) -> dict:
    t0 = time.perf_counter()
    if input_path.suffix.lower() not in (".docx", ".docm"):
        raise ValueError("Only .docx / .docm are supported. Convert .doc to .docx in Word.")

    if margins_mm is None and margin_preset not in MARGIN_PRESETS:
        raise ValueError(f"Unknown margin preset: {margin_preset}")

    mm = dict(margins_mm if margins_mm is not None else MARGIN_PRESETS[margin_preset])

    if backup:
        bak = input_path.with_suffix(input_path.suffix + ".gost_backup")
        shutil.copy2(input_path, bak)

    doc = Document(str(input_path))
    centered_ids = _snapshot_centered_paragraph_ids(doc)

    _apply_section_margins(doc, mm)
    _style_document_defaults(doc)

    if page_number == "footer":
        _setup_footer_page_numbers(doc)
    else:
        _setup_header_page_numbers(doc)

    eff_h = _effective_heading_count_for_toc(doc, centered_ids)
    want_toc = insert_toc if insert_toc is not None else eff_h > toc_min_headings
    toc_field_id: int | None = None
    if want_toc:
        toc_para = _insert_toc_paragraph(doc)
        toc_field_id = id(toc_para)
        centered_ids = _snapshot_centered_paragraph_ids(doc)

    if infer_captions:
        _ensure_table_captions(doc)
        _ensure_figure_captions(doc)
        centered_ids = _snapshot_centered_paragraph_ids(doc)

    if renumber:
        _renumber_captions(doc)

    paras = list(_iter_all_paragraphs(doc))
    total = len(paras)
    for i, p in enumerate(paras):
        if progress and total > 40 and i % 20 == 0:
            print(f"Formatting… {i}/{total}", file=sys.stderr)
        _process_paragraph(
            p,
            centered_ids=centered_ids,
            preserve_center=preserve_center,
            lab_heading_heuristic=lab_heading_heuristic,
            toc_field_id=toc_field_id,
        )

    _center_tables(doc)
    _center_image_paragraphs(doc)
    _apply_tnr_to_table_cells(doc)
    _apply_tnr_to_body_drawing_paragraphs(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    elapsed = time.perf_counter() - t0
    return {"output": str(output_path), "seconds": elapsed, "paragraphs": total}


def default_output_path(input_path: Path) -> Path:
    stem = input_path.stem
    return input_path.parent / f"ГОСТформат{stem}.docx"


def main() -> None:
    ap = argparse.ArgumentParser(
        description="Format Word .docx: типовая методичка / ГОСТ Р 7.0.11-2008 (вёрстка через python-docx)."
    )
    ap.add_argument("input", type=Path, help="Path to input .docx")
    ap.add_argument("-o", "--output", type=Path, default=None, help="Output path (default: ГОСТформат<name>.docx)")
    ap.add_argument(
        "--backup",
        action="store_true",
        help="Сохранить копию исходника рядом с расширением .gost_backup",
    )
    ap.add_argument("--no-renumber", action="store_true", help="Skip figure/table sequential renumbering")
    g = ap.add_mutually_exclusive_group()
    g.add_argument("--toc", dest="toc_mode", action="store_const", const="on", default=None, help="Always insert TOC")
    g.add_argument("--no-toc", dest="toc_mode", action="store_const", const="off", help="Never insert TOC")
    ap.set_defaults(toc_mode=None)
    ap.add_argument(
        "--page-number",
        choices=("header", "footer"),
        default="header",
        help="Положение номера страницы: верх (типично для сводок) или низ (прежний вариант skill)",
    )
    ap.add_argument(
        "--margin-preset",
        choices=tuple(MARGIN_PRESETS.keys()),
        default="methodology",
        help="methodology: 25/10/20/20 мм; legacy: 30/15/20/20 мм",
    )
    ap.add_argument(
        "--margins-mm",
        nargs=4,
        type=int,
        default=None,
        metavar=("L", "R", "T", "B"),
        help="Явные поля в мм: лево право верх низ (перекрывает --margin-preset)",
    )
    ap.add_argument(
        "--no-preserve-center",
        action="store_true",
        help="Не сохранять исходное центрирование (всё как основной текст)",
    )
    ap.add_argument(
        "--no-lab-headings",
        action="store_true",
        help="Отключить эвристику заголовков разделов (Цель, Задание, …)",
    )
    ap.add_argument(
        "--no-infer-captions",
        action="store_true",
        help="Не добавлять подписи к рисункам и таблицам по контексту",
    )
    ap.add_argument("--progress", action="store_true", help="Print progress to stderr for large files")
    args = ap.parse_args()
    out = args.output or default_output_path(args.input)
    insert_toc: bool | None
    if args.toc_mode == "on":
        insert_toc = True
    elif args.toc_mode == "off":
        insert_toc = False
    else:
        insert_toc = None
    mm_override: dict[str, int] | None = None
    if args.margins_mm is not None:
        mm_override = {
            "left": args.margins_mm[0],
            "right": args.margins_mm[1],
            "top": args.margins_mm[2],
            "bottom": args.margins_mm[3],
        }
    info = format_docx(
        args.input,
        out,
        backup=args.backup,
        insert_toc=insert_toc,
        renumber=not args.no_renumber,
        progress=args.progress,
        margin_preset=args.margin_preset,
        margins_mm=mm_override,
        page_number=args.page_number,
        preserve_center=not args.no_preserve_center,
        lab_heading_heuristic=not args.no_lab_headings,
        infer_captions=not args.no_infer_captions,
    )
    print("✅ Оформление применено (типовые параметры методичек / ГОСТ Р 7.0.11-2008).")
    print("   Уточняйте поля и колонтитулы по методичке вуза; см. SKILL.md и reference.md.")
    print(f"   Выходной файл: {info['output']}")
    print(f"   Абзацев обработано: {info['paragraphs']}, время: {info['seconds']:.2f} с")


if __name__ == "__main__":
    main()
